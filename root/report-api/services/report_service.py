from docx import Document
from docx.shared import Inches
from datetime import datetime
import os
import tempfile

class ReportService:
    def __init__(self):
        pass
    
    def generate_report(self, patient, medical_data):
        # Criar um novo documento Word
        doc = Document()
        
        # Título do documento
        title = doc.add_heading('Relatório Médico - Alzheimer', 0)
        title.alignment = 1  # Centralizar
        
        # Dados do paciente
        doc.add_heading('Dados do Paciente', level=1)
        patient_table = doc.add_table(rows=8, cols=2)
        patient_table.style = 'Table Grid'
        
        patient_data = [
            ('Nome:', patient.get('name', '')),
            ('Email:', patient.get('email', '')),
            ('Data de Nascimento:', patient.get('birthDate', '')),
            ('Gênero:', patient.get('gender', '')),
            ('Estado:', patient.get('state', '')),
            ('Etnia:', patient.get('ethnicity', '')),
            ('Nível Educacional:', patient.get('educationLevel', '')),
            ('Telefone:', patient.get('phoneNumber', ''))
        ]
        
        for i, (label, value) in enumerate(patient_data):
            patient_table.cell(i, 0).text = label
            patient_table.cell(i, 1).text = str(value)
        
        # Dados médicos
        doc.add_heading('Dados Médicos', level=1)
        doc.add_paragraph(f"Data do Exame: {medical_data.get('dateExam', '')}")
        
        # Fatores de estilo de vida
        doc.add_heading('Fatores de Estilo de Vida', level=2)
        lifestyle_table = doc.add_table(rows=6, cols=2)
        lifestyle_table.style = 'Table Grid'
        
        lifestyle_data = [
            ('Peso (kg):', medical_data.get('weight', '')),
            ('Altura (cm):', medical_data.get('height', '')),
            ('IMC:', medical_data.get('bmi', '')),
            ('Qualidade da Dieta (0-10):', medical_data.get('dietQuality', '')),
            ('Qualidade do Sono (0-10):', medical_data.get('sleepQuality', '')),
            ('Atividade Física:', 'Sim' if medical_data.get('physicalActivity') else 'Não')
        ]
        
        for i, (label, value) in enumerate(lifestyle_data):
            lifestyle_table.cell(i, 0).text = label
            lifestyle_table.cell(i, 1).text = str(value)
        
        # Histórico médico
        doc.add_heading('Histórico Médico', level=2)
        history_table = doc.add_table(rows=6, cols=2)
        history_table.style = 'Table Grid'
        
        history_data = [
            ('Histórico Familiar de Alzheimer:', 'Sim' if medical_data.get('familyHistory') else 'Não'),
            ('Doença Cardiovascular:', 'Sim' if medical_data.get('cardiovascularDisease') else 'Não'),
            ('Diabetes:', 'Sim' if medical_data.get('diabetes') else 'Não'),
            ('Depressão:', 'Sim' if medical_data.get('depression') else 'Não'),
            ('Traumatismo Craniano:', 'Sim' if medical_data.get('headTrauma') else 'Não'),
            ('Hipertensão:', 'Sim' if medical_data.get('hypertension') else 'Não')
        ]
        
        for i, (label, value) in enumerate(history_data):
            history_table.cell(i, 0).text = label
            history_table.cell(i, 1).text = str(value)
        
        # Avaliações cognitivas
        doc.add_heading('Avaliações Cognitivas', level=2)
        cognitive_table = doc.add_table(rows=5, cols=2)
        cognitive_table.style = 'Table Grid'
        
        cognitive_data = [
            ('MMSE (0-30):', medical_data.get('mmse', '')),
            ('ADL (0-10):', medical_data.get('adl', '')),
            ('Avaliação Funcional (0-10):', medical_data.get('functionalAssessment', '')),
            ('Queixas de Memória:', 'Sim' if medical_data.get('memoryComplaints') else 'Não'),
            ('Problemas Comportamentais:', 'Sim' if medical_data.get('behavioralProblems') else 'Não')
        ]
        
        for i, (label, value) in enumerate(cognitive_data):
            cognitive_table.cell(i, 0).text = label
            cognitive_table.cell(i, 1).text = str(value)
        
        # Sintomas
        doc.add_heading('Sintomas', level=2)
        symptoms_table = doc.add_table(rows=5, cols=2)
        symptoms_table.style = 'Table Grid'
        
        symptoms_data = [
            ('Confusão:', 'Sim' if medical_data.get('confusion') else 'Não'),
            ('Desorientação:', 'Sim' if medical_data.get('disorientation') else 'Não'),
            ('Esquecimento:', 'Sim' if medical_data.get('forgetfulness') else 'Não'),
            ('Mudanças de Personalidade:', 'Sim' if medical_data.get('personalityChanges') else 'Não'),
            ('Dificuldade em Completar Tarefas:', 'Sim' if medical_data.get('difficultyCompletingTasks') else 'Não')
        ]
        
        for i, (label, value) in enumerate(symptoms_data):
            symptoms_table.cell(i, 0).text = label
            symptoms_table.cell(i, 1).text = str(value)
        
        # Medições clínicas
        doc.add_heading('Medições Clínicas', level=2)
        clinical_table = doc.add_table(rows=6, cols=2)
        clinical_table.style = 'Table Grid'
        
        clinical_data = [
            ('Colesterol Total:', medical_data.get('cholesterolTotal', '')),
            ('Colesterol LDL:', medical_data.get('cholesterolLdl', '')),
            ('Colesterol HDL:', medical_data.get('cholesterolHdl', '')),
            ('Triglicerídeos:', medical_data.get('cholesterolTriglycerides', '')),
            ('Pressão Sistólica:', medical_data.get('systolicBP', '')),
            ('Pressão Diastólica:', medical_data.get('diastolicBP', ''))
        ]
        
        for i, (label, value) in enumerate(clinical_data):
            clinical_table.cell(i, 0).text = label
            clinical_table.cell(i, 1).text = str(value)
        
        # Salvar o documento em um arquivo temporário
        temp_dir = tempfile.gettempdir()
        filename = f"relatorio_{patient.get('name', 'paciente').replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(temp_dir, filename)
        
        doc.save(file_path)
        
        return file_path
